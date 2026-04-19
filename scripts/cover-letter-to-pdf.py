"""Convert cover letter .txt files to .docx then .pdf.

Usage: python scripts/cover-letter-to-pdf.py <slug> [<slug> ...]
Where slug is the filename stem in output/cover-letters/ (e.g., aledade-enterprise-data-architect).

Each .txt is split on blank lines into blocks: first block = title (heading),
remaining blocks = paragraphs preserving internal line breaks.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert

ROOT = Path(__file__).resolve().parent.parent
COVER_DIR = ROOT / "output" / "cover-letters"


def txt_to_docx(txt_path: Path, docx_path: Path) -> None:
    text = txt_path.read_text(encoding="utf-8").strip()
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if not blocks:
        raise ValueError(f"Empty cover letter: {txt_path}")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = blocks[0]
    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(13)

    for block in blocks[1:]:
        p = doc.add_paragraph()
        for i, line in enumerate(block.splitlines()):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)
        p.paragraph_format.space_after = Pt(8)

    doc.save(str(docx_path))


def main(slugs: list[str]) -> int:
    for slug in slugs:
        txt = COVER_DIR / f"{slug}.txt"
        docx = COVER_DIR / f"{slug}.docx"
        pdf = COVER_DIR / f"{slug}.pdf"

        if not txt.exists():
            print(f"[skip] {txt} not found", file=sys.stderr)
            continue

        print(f"[txt -> docx] {slug}")
        txt_to_docx(txt, docx)

        print(f"[docx -> pdf] {slug}")
        convert(str(docx), str(pdf))

        if not pdf.exists():
            print(f"[fail] PDF not produced: {pdf}", file=sys.stderr)
            return 1
        size = pdf.stat().st_size
        print(f"[ok]   {pdf.name} ({size:,} bytes)")

    return 0


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/cover-letter-to-pdf.py <slug> [<slug> ...]", file=sys.stderr)
        sys.exit(2)
    sys.exit(main(sys.argv[1:]))
