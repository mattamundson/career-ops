"""cv_docx_tailor — tailor the master CV DOCX + convert to PDF.

The master DOCX at output/Matt_Amundson_TOP_2026.docx is Matt's
professionally-designed resume. This helper:

1. Loads the master DOCX.
2. Patches the headline paragraph (index 1) and summary paragraph (index 4)
   with role-tailored text while preserving all formatting (font, size,
   alignment inherited from the original runs).
3. Saves a tailored copy and converts it to PDF via docx2pdf (Word COM
   automation on Windows).

Usage (CLI):
    python -m scripts.lib.cv_docx_tailor \
        --headline "Data Leader | Snowflake / dbt / Airflow" \
        --summary  "Data & analytics leader with 10+ years..." \
        --out      output/cv-acme-data-architect

    # Produces:
    #   output/cv-acme-data-architect.docx
    #   output/cv-acme-data-architect.pdf

Usage (as a library):
    from scripts.lib.cv_docx_tailor import tailor, tailor_and_pdf
    tailor_and_pdf(
        headline="...", summary="...",
        out_stem=Path("output/cv-foo"),
    )

The master DOCX is NEVER written to — we always write to a fresh copy.
"""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path
from typing import Optional

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
MASTER_DOCX = ROOT / "output" / "Matt_Amundson_TOP_2026.docx"

# Paragraph indices in the master DOCX that we rewrite per role.
# Determined by manual inspection of the master (see scripts/cv-docx-to-pdf.mjs).
HEADLINE_INDEX = 1
SUMMARY_INDEX = 4


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Rewrite paragraph text in-place, preserving the formatting of the
    first run. All subsequent runs are removed so we don't split style mid-line.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]
    first_run.text = new_text

    # Remove trailing runs — their text is now stale.
    for run in paragraph.runs[1:]:
        run.text = ""


def tailor(
    headline: Optional[str],
    summary: Optional[str],
    out_docx: Path,
    master: Path = MASTER_DOCX,
) -> Path:
    """Copy the master DOCX to out_docx and patch the headline + summary.

    Returns the path to the written DOCX.
    """
    if not master.exists():
        raise FileNotFoundError(f"Master DOCX not found: {master}")

    out_docx = Path(out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(master, out_docx)

    doc = Document(str(out_docx))
    paragraphs = doc.paragraphs

    if len(paragraphs) <= max(HEADLINE_INDEX, SUMMARY_INDEX):
        raise ValueError(
            f"Master DOCX has {len(paragraphs)} paragraphs; expected at least "
            f"{max(HEADLINE_INDEX, SUMMARY_INDEX) + 1}. Did the template change?"
        )

    if headline is not None:
        _replace_paragraph_text(paragraphs[HEADLINE_INDEX], headline)

    if summary is not None:
        _replace_paragraph_text(paragraphs[SUMMARY_INDEX], summary)

    doc.save(str(out_docx))
    return out_docx


def tailor_and_pdf(
    headline: Optional[str],
    summary: Optional[str],
    out_stem: Path,
    master: Path = MASTER_DOCX,
) -> tuple[Path, Path]:
    """Tailor the DOCX and produce a PDF alongside it.

    out_stem may include .docx or .pdf — the suffix is stripped and both
    extensions are written. Returns (docx_path, pdf_path).
    """
    from docx2pdf import convert

    out_stem = Path(out_stem)
    if out_stem.suffix.lower() in {".docx", ".pdf"}:
        out_stem = out_stem.with_suffix("")

    docx_path = out_stem.with_suffix(".docx")
    pdf_path = out_stem.with_suffix(".pdf")

    tailor(headline, summary, docx_path, master=master)
    convert(str(docx_path), str(pdf_path))

    if not pdf_path.exists():
        raise RuntimeError(f"docx2pdf ran but PDF was not produced: {pdf_path}")

    return docx_path, pdf_path


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--headline", type=str, default=None,
                    help="Role-tailored headline line (paragraph 1 of master)")
    ap.add_argument("--summary", type=str, default=None,
                    help="Role-tailored professional summary (paragraph 4)")
    ap.add_argument("--out", type=Path, required=True,
                    help="Output path stem (e.g., output/cv-acme). "
                         ".docx and .pdf are appended.")
    ap.add_argument("--docx-only", action="store_true",
                    help="Write the tailored DOCX but skip PDF conversion.")
    ap.add_argument("--master", type=Path, default=MASTER_DOCX,
                    help=f"Override master DOCX (default: {MASTER_DOCX})")
    args = ap.parse_args(argv)

    if args.headline is None and args.summary is None:
        print("[cv-docx-tailor] no --headline or --summary provided; "
              "will still copy master unchanged.", file=sys.stderr)

    if args.docx_only:
        stem = Path(args.out)
        if stem.suffix.lower() in {".docx", ".pdf"}:
            stem = stem.with_suffix("")
        docx_path = stem.with_suffix(".docx")
        tailor(args.headline, args.summary, docx_path, master=args.master)
        print(f"[cv-docx-tailor] wrote {docx_path}")
        return 0

    try:
        docx_path, pdf_path = tailor_and_pdf(
            args.headline, args.summary, args.out, master=args.master
        )
    except Exception as err:
        print(f"[cv-docx-tailor] FAILED: {err}", file=sys.stderr)
        return 1

    size = pdf_path.stat().st_size
    print(f"[cv-docx-tailor] wrote {docx_path}")
    print(f"[cv-docx-tailor] wrote {pdf_path} ({size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
